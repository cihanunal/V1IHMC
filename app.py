import datetime
import io
import re
from typing import Dict, List, Optional, Tuple

import pandas as pd
import streamlit as st
import xlsxwriter

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

try:
    import gspread
    from google.oauth2.service_account import Credentials
except Exception:
    gspread = None
    Credentials = None


# =========================
# AYARLAR
# =========================

DEFAULT_KATILIMCI_SHEET_ID = "1b6K9svdyT_RrgGCFsGYO5nrmPPYw6pkvO-BTzfwNT9k"
DEFAULT_PROGRAM_SHEET_ID = "1Polxg5n-J0VueJifvjlgoGvXITVvnBbJpgSjHJ6imYY"

GOOGLE_SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive.readonly",
]

PROGRAM_PLAN_INDEX = 0
PROGRAM_BILDIRILER_SHEET = "Bildiriler"
PROGRAM_OZEL_SHEET = "Ozel_Etkinlikler"
PROGRAM_MOD_SHEET = "Moderatorler"


# =========================
# GENEL YARDIMCILAR
# =========================

def safe_str(value) -> str:
    if pd.isna(value):
        return ""
    return str(value).strip()


def turkce_buyuk(metin):
    if not isinstance(metin, str) or not metin:
        return ""
    return (
        metin.replace("i", "İ")
        .replace("ı", "I")
        .replace("ş", "Ş")
        .replace("ğ", "Ğ")
        .replace("ü", "Ü")
        .replace("ö", "Ö")
        .replace("ç", "Ç")
        .upper()
    )


def temiz_metin(metin):
    if not isinstance(metin, str) or not metin:
        return ""
    m = (
        metin.upper()
        .replace("İ", "I")
        .replace("Ğ", "G")
        .replace("Ü", "U")
        .replace("Ş", "S")
        .replace("Ö", "O")
        .replace("Ç", "C")
    )
    m = re.sub(r"[^A-Z0-9 ]", "", m)
    return re.sub(r"\s+", " ", m).strip()


def super_temiz(metin):
    if not isinstance(metin, str) or not metin:
        return ""
    m = turkce_buyuk(metin)
    return (
        m.replace("Ö", "O")
        .replace("Ü", "U")
        .replace("Ş", "S")
        .replace("Ç", "C")
        .replace("Ğ", "G")
        .replace("İ", "I")
    )


def temiz_isim(isim):
    if not isinstance(isim, str) or str(isim).strip() == "":
        return ""
    i = (
        str(isim)
        .upper()
        .replace("İ", "I")
        .replace("Ğ", "G")
        .replace("Ü", "U")
        .replace("Ş", "S")
        .replace("Ö", "O")
        .replace("Ç", "C")
    )
    unvan_paternleri = [
        r"^PROF\.?\s*DR\.?",
        r"^DOC\.?\s*DR\.?",
        r"^DOÇ\.?\s*DR\.?",
        r"^DR\.?\s*OGR\.?\s*UYESI",
        r"^DR\.?\s*ÖĞR\.?\s*ÜYESI",
        r"^DR\.?",
        r"^UZM\.?",
        r"^ARAS\.?\s*GOR\.?",
        r"^ARŞ\.?\s*GÖR\.?",
        r"^ARS\.?\s*GOR\.?",
        r"^ASSOC\.?\s*PROF\.?",
        r"^ASSIST\.?\s*PROF\.?",
        r"^PROF\.?",
    ]
    for patern in unvan_paternleri:
        i = re.sub(patern, "", i).strip()
    return re.sub(r"\s+", " ", i).strip()


def unvan_bul(row, columns, raw_val):
    for col in columns:
        if "unvan" in str(col).lower():
            u = str(row[col]).strip()
            if u and u.lower() != "nan":
                return u
    raw_upper = str(raw_val).upper()
    if "ASSOC. PROF" in raw_upper or "ASSOC PROF" in raw_upper:
        return "Assoc. Prof."
    if "ASSIST. PROF" in raw_upper or "ASSIST PROF" in raw_upper:
        return "Assist. Prof."
    if "PROF" in raw_upper:
        return "Prof. Dr."
    if "DOÇ" in raw_upper or "DOC" in raw_upper:
        return "Doç. Dr."
    if "ÖĞR" in raw_upper or "OGR" in raw_upper:
        return "Dr. Öğr. Üyesi"
    if "DR." in raw_upper or "DR " in raw_upper:
        return "Dr."
    if "UZM" in raw_upper:
        return "Uzm."
    if "ARŞ" in raw_upper or "ARS" in raw_upper:
        return "Arş. Gör."
    return ""


def find_col(df: pd.DataFrame, candidates: List[str], contains: Optional[List[str]] = None) -> Optional[str]:
    if df is None or df.empty:
        return None
    cols = list(df.columns)
    normalized = {temiz_metin(str(c)): c for c in cols}
    for cand in candidates:
        key = temiz_metin(cand)
        if key in normalized:
            return normalized[key]
    if contains:
        for col in cols:
            col_low = str(col).lower()
            if all(x.lower() in col_low for x in contains):
                return col
    return None


def parse_zaman(metin):
    match = re.search(r"(\d{2}\.\d{2}\.\d{4}).*?(\d{2}[:.]\d{2})", str(metin))
    if match:
        date_str = match.group(1)
        time_str = match.group(2).replace(".", ":")
        try:
            return datetime.datetime.strptime(f"{date_str} {time_str}", "%d.%m.%Y %H:%M")
        except Exception:
            pass
    return datetime.datetime(2099, 1, 1)


def gun_key(gun_ve_saat: str) -> str:
    text = str(gun_ve_saat).strip()
    if " | " in text:
        return text.split(" | ")[0].strip()
    m = re.search(r"\d{2}\.\d{2}\.\d{4}\s+\S+", text)
    if m:
        return m.group(0)
    return text


def clean_salon(sal):
    if isinstance(sal, pd.Timestamp) or hasattr(sal, "month"):
        if sal.month == 8 and sal.day == 30:
            return "30 Ağustos Salonu"

    n = str(sal).upper().strip()
    if "30 A" in n or "08-30" in n:
        return "30 Ağustos Salonu"
    if "MALAZGIRT" in n or "MALAZGİRT" in n:
        return "Malazgirt Salonu"
    if "ELIF" in n or "ELİF" in n:
        return "Doç. Dr. Elif Kaya Salonu"
    if "AHMET" in n or "EKIZER" in n or "EKİZER" in n:
        return "Dr. Ahmet Ekizer Salonu"
    if "TEAMS 1" in n or "ZOOM 1" in n or "ODA 1" in n:
        return "Teams Oda 1"
    if "TEAMS 2" in n or "ZOOM 2" in n or "ODA 2" in n:
        return "Teams Oda 2"
    if "TEAMS 3" in n or "ZOOM 3" in n or "ODA 3" in n:
        return "Teams Oda 3"
    return str(sal).strip()


def get_kisi_renk_emoji(k_data):
    odeme = str(k_data.get("Odeme", "")).lower()
    gorev = str(k_data.get("Gorev", "")).lower()
    bildiri_sayisi = len(k_data.get("Bildiriler", []))
    if "muaf" in odeme or "görevli" in gorev or "davetli" in gorev:
        return "🟣"
    if "indirim" in odeme or "öğrenci" in odeme or "ogrenci" in odeme:
        return "🟠"
    if "bekleniyor" in odeme:
        return "🟡"
    if "yok" in odeme:
        return "⚪"
    if bildiri_sayisi == 0:
        return "🔵"
    return "🟢"


def get_bildiri_renk_emoji(b_data):
    return "🟢" if b_data.get("Odeme") == "Evet" else "🟡"


def get_program_info(b_isim, df_prog):
    if df_prog is None or df_prog.empty:
        return None
    b_isim_temiz = super_temiz(b_isim)
    for _, row in df_prog.iterrows():
        row_vals = [str(x) for x in row.values if pd.notna(x)]
        row_str = super_temiz(" ".join(row_vals))
        if b_isim_temiz and b_isim_temiz in row_str:
            parts = []
            for key in ["Gun_ve_Saat", "Tarih", "Saat", "Salon", "Oturum_ID", "Sunum_Tipi"]:
                if key in df_prog.columns and safe_str(row.get(key)):
                    parts.append(f"{key}: {safe_str(row.get(key))}")
            return " - ".join(parts) if parts else " | ".join(row_vals[:4])
    return "Henüz program tablosunda yeri belli değil."


# =========================
# GOOGLE SHEETS OKUMA/YAZMA
# =========================

def has_service_account() -> bool:
    try:
        return bool(st.secrets.get("gcp_service_account"))
    except Exception:
        return False


@st.cache_resource(show_spinner=False)
def get_gspread_client():
    if gspread is None or Credentials is None:
        return None
    if not has_service_account():
        return None
    info = dict(st.secrets["gcp_service_account"])
    creds = Credentials.from_service_account_info(info, scopes=GOOGLE_SCOPES)
    return gspread.authorize(creds)


def worksheet_to_df(ws) -> pd.DataFrame:
    values = ws.get_all_values()
    if not values:
        return pd.DataFrame()
    headers = [str(h).strip() for h in values[0]]
    rows = values[1:]
    df = pd.DataFrame(rows, columns=headers).fillna("")
    df["_row"] = list(range(2, len(df) + 2))
    return df


def read_ws_by_name(book, sheet_name: str) -> pd.DataFrame:
    try:
        return worksheet_to_df(book.worksheet(sheet_name))
    except Exception:
        return pd.DataFrame()


def read_ws_by_index(book, index: int) -> pd.DataFrame:
    try:
        return worksheet_to_df(book.get_worksheet(index))
    except Exception:
        return pd.DataFrame()


def public_excel_url(sheet_id: str) -> str:
    return f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=xlsx"


def read_public_sheet(sheet_id: str, sheet_name=None) -> pd.DataFrame:
    xls = pd.ExcelFile(public_excel_url(sheet_id))
    df = pd.read_excel(xls, sheet_name=sheet_name, dtype=str).fillna("")
    if isinstance(df, dict):
        return pd.DataFrame()
    df.columns = [str(c).strip() for c in df.columns]
    df["_row"] = list(range(2, len(df) + 2))
    return df


@st.cache_data(ttl=600, show_spinner=False)
def load_all_data(katilimci_sheet_id: str, program_sheet_id: str, writable_mode: bool) -> Dict[str, pd.DataFrame]:
    data = {}
    client = get_gspread_client() if writable_mode else None

    if client:
        kat_book = client.open_by_key(katilimci_sheet_id)
        prog_book = client.open_by_key(program_sheet_id)

        data["program_plan"] = read_ws_by_index(prog_book, PROGRAM_PLAN_INDEX)
        data["program_bildiriler"] = read_ws_by_name(prog_book, PROGRAM_BILDIRILER_SHEET)
        data["ozel"] = read_ws_by_name(prog_book, PROGRAM_OZEL_SHEET)
        data["mod"] = read_ws_by_name(prog_book, PROGRAM_MOD_SHEET)

        data["bildiri_liste"] = read_ws_by_name(kat_book, "kabul edilen bildiriler")
        data["odeme"] = read_ws_by_name(kat_book, "ödeme durumu")
        data["bilgi"] = read_ws_by_name(kat_book, "katılımcı bilgileri")
        data["anket"] = read_ws_by_name(kat_book, "katılımcıların ankete cevapları")
        data["danisma"] = read_ws_by_name(kat_book, "Kongre Bilimsel Danışma Kurulu")
        data["duzenleme"] = read_ws_by_name(kat_book, "Kongre Düzenleme Kurulu")
        data["bildiri_detay"] = read_ws_by_name(kat_book, "bildirilerin detayları")
        data["yaka_ek"] = read_ws_by_name(kat_book, "yaka kartı ek liste")
        return data

    # Sadece okuma modu: herkese açık export linkleri ile çalışır.
    prog_xls = pd.ExcelFile(public_excel_url(program_sheet_id))
    kat_xls = pd.ExcelFile(public_excel_url(katilimci_sheet_id))

    def read_from_xls(xls, sheet_name):
        try:
            df = pd.read_excel(xls, sheet_name=sheet_name, dtype=str).fillna("")
            df.columns = [str(c).strip() for c in df.columns]
            df["_row"] = list(range(2, len(df) + 2))
            return df
        except Exception:
            return pd.DataFrame()

    try:
        df_plan = pd.read_excel(prog_xls, sheet_name=0, dtype=str).fillna("")
        df_plan.columns = [str(c).strip() for c in df_plan.columns]
        df_plan["_row"] = list(range(2, len(df_plan) + 2))
        data["program_plan"] = df_plan
    except Exception:
        data["program_plan"] = pd.DataFrame()

    data["program_bildiriler"] = read_from_xls(prog_xls, PROGRAM_BILDIRILER_SHEET)
    data["ozel"] = read_from_xls(prog_xls, PROGRAM_OZEL_SHEET)
    data["mod"] = read_from_xls(prog_xls, PROGRAM_MOD_SHEET)
    data["bildiri_liste"] = read_from_xls(kat_xls, "kabul edilen bildiriler")
    data["odeme"] = read_from_xls(kat_xls, "ödeme durumu")
    data["bilgi"] = read_from_xls(kat_xls, "katılımcı bilgileri")
    data["anket"] = read_from_xls(kat_xls, "katılımcıların ankete cevapları")
    data["danisma"] = read_from_xls(kat_xls, "Kongre Bilimsel Danışma Kurulu")
    data["duzenleme"] = read_from_xls(kat_xls, "Kongre Düzenleme Kurulu")
    data["bildiri_detay"] = read_from_xls(kat_xls, "bildirilerin detayları")
    data["yaka_ek"] = read_from_xls(kat_xls, "yaka kartı ek liste")
    return data


def get_worksheet(sheet_id: str, sheet_name_or_index):
    client = get_gspread_client()
    if not client:
        raise RuntimeError("Google Sheets yazma için Streamlit secrets içine gcp_service_account eklenmeli.")
    book = client.open_by_key(sheet_id)
    if isinstance(sheet_name_or_index, int):
        return book.get_worksheet(sheet_name_or_index)
    return book.worksheet(sheet_name_or_index)


def update_sheet_row(sheet_id: str, sheet_name_or_index, row_number: int, updates: Dict[str, str]):
    ws = get_worksheet(sheet_id, sheet_name_or_index)
    headers = [h.strip() for h in ws.row_values(1)]
    header_map = {h: i + 1 for i, h in enumerate(headers)}
    missing = [col for col in updates if col not in header_map]
    if missing:
        raise ValueError(f"Sheet içinde şu kolonlar bulunamadı: {', '.join(missing)}")

    payload = []
    for col_name, value in updates.items():
        a1 = gspread.utils.rowcol_to_a1(row_number, header_map[col_name])
        payload.append({"range": a1, "values": [[value]]})
    ws.batch_update(payload, value_input_option="USER_ENTERED")


def find_row_by_value(sheet_id: str, sheet_name_or_index, col_name: str, target_value: str) -> Optional[int]:
    ws = get_worksheet(sheet_id, sheet_name_or_index)
    values = ws.get_all_values()
    if not values:
        return None
    headers = [h.strip() for h in values[0]]
    if col_name not in headers:
        return None
    col_idx = headers.index(col_name)
    target = temiz_metin(target_value)
    for i, row in enumerate(values[1:], start=2):
        val = row[col_idx] if col_idx < len(row) else ""
        if temiz_metin(val) == target:
            return i
    return None


# =========================
# VERİ MODELİ
# =========================

def build_database(raw_data: Dict[str, pd.DataFrame]) -> Tuple[Dict, Dict]:
    bildiriler = {}
    katilimcilar = {}
    detay_map = {}

    df_detay = raw_data.get("bildiri_detay", pd.DataFrame())
    for _, row in df_detay.iterrows():
        b_adi = ""
        for col in df_detay.columns:
            if "bildiri" in str(col).lower() and ("ad" in str(col).lower() or "ism" in str(col).lower()):
                b_adi = str(row[col]).strip()
                break
        if b_adi:
            b_key = temiz_metin(b_adi)
            detay_map[b_key] = {
                "Konu": str(row.get("Konu", "")),
                "Butik": "Evet" if "evet" in str(row.get("Butik Bildiri", "")).lower() else "",
            }

    df_bildiri = raw_data.get("program_bildiriler")
    if df_bildiri is None or df_bildiri.empty:
        df_bildiri = raw_data.get("bildiri_liste", pd.DataFrame())

    b_col = find_col(df_bildiri, ["Bildiri Ismi", "Bildiri_Adi", "Bildiri Adı"], contains=["bildiri"])
    sun_col = find_col(df_bildiri, ["Sunan Yazar", "sunucu", "Sunucu"], contains=["sunan"])

    for _, row in df_bildiri.iterrows():
        b_adi_orijinal = str(row.get(b_col, "")).strip() if b_col else ""
        if b_adi_orijinal:
            b_key = temiz_metin(b_adi_orijinal)
            sunucu = temiz_isim(str(row.get(sun_col, ""))) if sun_col else ""
            if b_key not in bildiriler:
                bildiriler[b_key] = {
                    "Orijinal İsim": b_adi_orijinal,
                    "Konu": detay_map.get(b_key, {}).get("Konu", str(row.get("Konu", "-"))),
                    "Butik": detay_map.get(b_key, {}).get("Butik", ""),
                    "Kabul": "Evet",
                    "Yazarlar": [],
                    "Sunucu": sunucu,
                    "Odeme": "Hayır",
                }
            for i in range(1, 9):
                y = temiz_isim(str(row.get(f"Yazar {i}", "")))
                if y and y not in bildiriler[b_key]["Yazarlar"]:
                    bildiriler[b_key]["Yazarlar"].append(y)
            if sunucu and sunucu not in bildiriler[b_key]["Yazarlar"]:
                bildiriler[b_key]["Yazarlar"].append(sunucu)

    def kisi_islet(isim, unvan="", fiziki=False, gorev="Yok", odeme="Ödeme Bekleniyor"):
        t = temiz_isim(isim)
        if not t:
            return
        if t not in katilimcilar:
            katilimcilar[t] = {
                "Orijinal İsim": turkce_buyuk(isim),
                "Unvan": unvan,
                "E-Posta": "",
                "Telefon": "",
                "Kurum": "",
                "Gorev": gorev,
                "Odeme": odeme,
                "Tur": "Fiziki" if fiziki else "Belirtilmedi",
                "Niyet": "Hayır",
                "Gunler": {"6": "", "7": "", "8": ""},
                "Etkinlikler": "",
                "Bildiriler": [],
            }
        if unvan:
            katilimcilar[t]["Unvan"] = unvan
        if gorev != "Yok":
            katilimcilar[t]["Gorev"] = gorev
        if odeme != "Ödeme Bekleniyor":
            katilimcilar[t]["Odeme"] = odeme

    for sheet_key, gorev in [
        ("danisma", "Görevli - Bilimsel Danışma Kurulu"),
        ("duzenleme", "Görevli - Kongre Düzenleme Kurulu"),
    ]:
        df = raw_data.get(sheet_key, pd.DataFrame())
        for _, r in df.iterrows():
            isim = str(r.iloc[0]).strip() if len(r) else ""
            unv = unvan_bul(r, df.columns, isim)
            kisi_islet(isim, unvan=unv, gorev=gorev, odeme="Görevli/Davetli (Muaf)", fiziki=True)

    df_yaka = raw_data.get("yaka_ek", pd.DataFrame())
    for _, r in df_yaka.iterrows():
        if len(r) >= 2:
            kisi_islet(
                str(r.iloc[1]).strip(),
                unvan=str(r.iloc[0]).strip(),
                gorev=str(r.iloc[2]).strip() if len(r) > 2 else "Davetli Katılımcı",
                odeme="Görevli/Davetli (Muaf)",
                fiziki=True,
            )

    df_odeme = raw_data.get("odeme", pd.DataFrame())
    for _, r in df_odeme.iterrows():
        if len(r) >= 2:
            kisi_islet(r.iloc[0], odeme=str(r.iloc[1]))

    df_bilgi = raw_data.get("bilgi", pd.DataFrame())
    for _, r in df_bilgi.iterrows():
        ad_soy = str(r.get("Adı Soyadı", r.iloc[0] if len(r) else "")).strip()
        t = temiz_isim(ad_soy)
        if t in katilimcilar:
            for c in df_bilgi.columns:
                cl = str(c).lower()
                val = str(r[c]).strip()
                if val and val.lower() != "nan":
                    if "mail" in cl or "posta" in cl:
                        katilimcilar[t]["E-Posta"] = val
                    elif "telefon" in cl or "gsm" in cl:
                        katilimcilar[t]["Telefon"] = val
                    elif "kurum" in cl or "üniversite" in cl:
                        katilimcilar[t]["Kurum"] = val

    df_anket = raw_data.get("anket", pd.DataFrame())
    for _, r in df_anket.iterrows():
        ad = str(r.get("Adı Soyadı", r.iloc[0] if len(r) else "")).strip()
        t = temiz_isim(ad)
        if not t:
            continue
        unv = unvan_bul(r, df_anket.columns, ad)
        kisi_islet(ad, unvan=unv)
        for c in df_anket.columns:
            cl = str(c).lower()
            val = str(r[c])
            if val.lower() == "nan" or not val:
                continue
            if "nasıl katılım" in cl:
                katilimcilar[t]["Tur"] = val
            elif "6 mayıs" in cl:
                katilimcilar[t]["Gunler"]["6"] = val
            elif "7 mayıs" in cl:
                katilimcilar[t]["Gunler"]["7"] = val
            elif "8 mayıs" in cl:
                katilimcilar[t]["Gunler"]["8"] = val
            elif "etkinlik" in cl:
                katilimcilar[t]["Etkinlikler"] += ", " + val
            elif "sunacaksanız" in cl and "katılımcı" not in val.lower():
                katilimcilar[t]["Niyet"] = "Evet"

    for b_key, b_v in bildiriler.items():
        bildiri_odendi = False
        for y in b_v["Yazarlar"]:
            if y not in katilimcilar:
                kisi_islet(y)
            if b_v["Orijinal İsim"] not in katilimcilar[y]["Bildiriler"]:
                katilimcilar[y]["Bildiriler"].append(b_v["Orijinal İsim"])
            odeme = str(katilimcilar[y]["Odeme"])
            if "Bekleniyor" not in odeme and "Yok" not in odeme:
                bildiri_odendi = True
        if bildiri_odendi:
            bildiriler[b_key]["Odeme"] = "Evet"

    for _, v in katilimcilar.items():
        if v["Odeme"] == "Ödeme Bekleniyor" and len(v["Bildiriler"]) == 0:
            v["Odeme"] = "Bildirisi Yok / Kayıt Yok"

    return bildiriler, katilimcilar


def build_matbaa_master(raw_data: Dict[str, pd.DataFrame]) -> pd.DataFrame:
    df_ayar = raw_data.get("program_plan", pd.DataFrame()).copy().fillna("-")
    df_bildiriler = raw_data.get("program_bildiriler", pd.DataFrame()).copy().fillna("-")
    if df_bildiriler.empty:
        df_bildiriler = raw_data.get("bildiri_liste", pd.DataFrame()).copy().fillna("-")

    for df in [df_ayar, df_bildiriler]:
        df.columns = [str(c).strip() for c in df.columns]

    yazar_sutunlari = [f"Yazar {i}" for i in range(1, 9)]
    for col in yazar_sutunlari:
        if col not in df_bildiriler.columns:
            df_bildiriler[col] = "-"

    isim_col = find_col(df_bildiriler, ["Bildiri Ismi", "Bildiri_Adi", "Bildiri Adı"], contains=["bildiri"])
    sunan_col = find_col(df_bildiriler, ["Sunan Yazar", "sunucu", "Sunucu"], contains=["sunan"])
    konu_col = find_col(df_bildiriler, ["Konu"], contains=["konu"])

    if isim_col is None:
        return pd.DataFrame()

    df_bildiriler["yazarlar"] = df_bildiriler.apply(
        lambda r: " - ".join([str(r[c]).strip() for c in yazar_sutunlari if str(r[c]).strip() not in ["-", ""]]),
        axis=1,
    )
    df_bildiriler["etkinlik_adi"] = df_bildiriler[isim_col].astype(str).str.strip()
    df_bildiriler["sunucu"] = df_bildiriler[sunan_col].astype(str).str.strip() if sunan_col else "-"
    df_bildiriler["konu"] = df_bildiriler[konu_col].astype(str).str.strip() if konu_col else "-"
    df_bildiriler = df_bildiriler.drop_duplicates(subset=["etkinlik_adi"])

    if "Bildiri_Adi" not in df_ayar.columns:
        possible = find_col(df_ayar, ["Bildiri_Adi", "Bildiri Adı", "Bildiri Ismi"], contains=["bildiri"])
        if possible:
            df_ayar["Bildiri_Adi"] = df_ayar[possible]
        else:
            return pd.DataFrame()

    df_ayar["Bildiri_Adi"] = df_ayar["Bildiri_Adi"].astype(str).str.strip()
    df_master = pd.merge(
        df_ayar,
        df_bildiriler[["etkinlik_adi", "yazarlar", "sunucu", "konu"]],
        left_on="Bildiri_Adi",
        right_on="etkinlik_adi",
        how="left",
    ).fillna("-")
    for col in ["Gun_ve_Saat", "Salon", "Oturum_ID", "Sunum_Tipi", "Bildiri_Adi", "sunucu", "yazarlar", "konu"]:
        if col not in df_master.columns:
            df_master[col] = "-"
    return df_master


# =========================
# MATBAA MOTORU
# =========================

def set_cell_bg(cell, hex_color):
    hex_color = hex_color.replace("#", "")
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_merged_row(table, text, bg_color, text_color=(0, 0, 0), bold=False, align=None):
    row = table.add_row()
    cell = row.cells[0]
    cell.merge(row.cells[1])
    cell.text = text
    set_cell_bg(cell, bg_color)
    for paragraph in cell.paragraphs:
        if align == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(*text_color)
            run.font.bold = bold
    return row


def build_program_dict(df_bildiriler):
    prog = {}
    for _, r in df_bildiriler.iterrows():
        b_adi = str(r["Bildiri_Adi"])
        sunucu = str(r["sunucu"])
        yazarlar = str(r["yazarlar"])
        konu = str(r["konu"])
        kz = str(r["Gun_ve_Saat"]).replace("Persembe", "Perşembe")
        ks = clean_salon(r["Salon"])
        sid = str(r["Oturum_ID"])
        if kz in ["", "-", "nan", "NaN"] or b_adi in ["", "-", "nan", "NaN"]:
            continue
        anahtar = (kz, ks)
        if anahtar not in prog:
            prog[anahtar] = []
        prog[anahtar].append({"b": b_adi, "y": yazarlar, "s": sunucu, "k": konu, "sid": sid})
    return prog


def moderator_atamalari_olustur(prog, df_mod, tip):
    mod_atamalari = {
        k: {"mod": "Daha Sonra İlan Edilecektir", "deg": "Daha Sonra İlan Edilecektir"} for k in prog.keys()
    }
    if df_mod is None or df_mod.empty:
        return mod_atamalari

    musait_oturumlar_mod = list(prog.keys())
    musait_oturumlar_deg = list(prog.keys())
    serbest_mods, serbest_degs = [], []

    for _, r in df_mod.iterrows():
        is_online_mod = str(r.get("Online", "-")).strip().lower() in [
            "evet",
            "e",
            "yes",
            "1",
            "var",
            "online",
            "true",
        ]
        if tip == "ONLINE" and not is_online_mod:
            continue
        if tip == "YUZYUZE" and is_online_mod:
            continue

        m_gun_saat = str(r.get("Gun_ve_Saat", "-")).strip().replace("Persembe", "Perşembe")
        m_salon_ham = str(r.get("Salon", "-")).strip()
        m_salon = clean_salon(m_salon_ham) if m_salon_ham != "-" else "-"
        m_oturum = str(r.get("Oturum_ID", "-")).strip()
        gorev_sutunu = str(r.get("Gorev", "-")).strip().lower()
        kurum = str(r.get("kurum", "-"))
        mod_metni = (
            f"{r.get('unvan_ad_soyad', '-')} ({r.get('kurum', '')})"
            if kurum not in ["-", "nan", ""]
            else str(r.get("unvan_ad_soyad", "-"))
        )

        is_deg = "deg" in gorev_sutunu
        role_key = "deg" if is_deg else "mod"
        hedef_musait_liste = musait_oturumlar_deg if is_deg else musait_oturumlar_mod
        hedef_serbest_liste = serbest_degs if is_deg else serbest_mods

        atandi_mi = False
        if m_gun_saat != "-" and m_salon != "-":
            hedef_tuple = (m_gun_saat, m_salon)
            if hedef_tuple in hedef_musait_liste:
                mod_atamalari[hedef_tuple][role_key] = mod_metni
                hedef_musait_liste.remove(hedef_tuple)
                atandi_mi = True

        if not atandi_mi and (m_gun_saat != "-" or m_salon != "-" or m_oturum != "-"):
            for sess in list(hedef_musait_liste):
                zaman_metni = sess[0]
                salon_metni = sess[1]
                oturum_id_metni = prog[sess][0]["sid"] if prog[sess] else "-"
                if (
                    (m_gun_saat == "-" or m_gun_saat in zaman_metni)
                    and (m_salon == "-" or m_salon == salon_metni)
                    and (m_oturum == "-" or m_oturum == oturum_id_metni)
                ):
                    mod_atamalari[sess][role_key] = mod_metni
                    hedef_musait_liste.remove(sess)
                    atandi_mi = True
                    break
        if not atandi_mi:
            hedef_serbest_liste.append(mod_metni)

    for mod_metni in serbest_mods:
        if musait_oturumlar_mod:
            mod_atamalari[musait_oturumlar_mod.pop(0)]["mod"] = mod_metni
    for deg_metni in serbest_degs:
        if musait_oturumlar_deg:
            mod_atamalari[musait_oturumlar_deg.pop(0)]["deg"] = deg_metni
    return mod_atamalari


def ozel_etkinlikleri_hazirla(df_ozel):
    soft_palette = ["#DDEBF7", "#FCE4D6", "#E2EFDA", "#FFF2CC", "#E6E6FA", "#F2F2F2"]
    gosterilecek = []
    if df_ozel is None or df_ozel.empty:
        return gosterilecek
    if not {"oturum_sirasi", "salon"}.issubset(set(df_ozel.columns)):
        return gosterilecek
    color_index = 0
    for (sira, sal), grup in df_ozel.groupby(["oturum_sirasi", "salon"], sort=False):
        ks = clean_salon(sal)
        renk_temasi = soft_palette[color_index % len(soft_palette)]
        color_index += 1
        ts = str(grup.iloc[0].get("tarih_saat", "-")).strip().replace("Persembe", "Perşembe")
        gosterilecek.append((sira, ks, grup, renk_temasi, ts))
    return gosterilecek


def gun_istatistikleri_olustur(prog):
    gun_istatistikleri = {}
    for (oturum_zaman, _sal), bilds in prog.items():
        t = oturum_zaman.split(" | ")[0] if " | " in oturum_zaman else oturum_zaman
        if t not in gun_istatistikleri:
            gun_istatistikleri[t] = {"oturum_sayisi": 0, "bildiri_sayisi": 0}
        gun_istatistikleri[t]["oturum_sayisi"] += 1
        gun_istatistikleri[t]["bildiri_sayisi"] += len(bilds)
    return gun_istatistikleri


def word_bas(df_bildiriler, df_ozel, df_mod, tip):
    prog = build_program_dict(df_bildiriler)
    gosterilecek_ozel_etkinlikler = ozel_etkinlikleri_hazirla(df_ozel)
    mod_atamalari = moderator_atamalari_olustur(prog, df_mod, tip)
    gun_istatistikleri = gun_istatistikleri_olustur(prog)

    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    ana_baslik = "11. IHMC FİZİKİ / FACE TO FACE PROGRAM" if tip == "YUZYUZE" else "11. IHMC ONLINE (DIGITAL) PROGRAM"
    h = doc.add_heading(ana_baslik, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    if len(gosterilecek_ozel_etkinlikler) > 0:
        t_ozel_baslik = doc.add_table(rows=0, cols=2)
        t_ozel_baslik.style = "Table Grid"
        add_merged_row(t_ozel_baslik, "ÖZEL ETKİNLİKLER", "000000", text_color=(255, 255, 255), bold=True, align="center")
        doc.add_paragraph()

        for _sira, ks, grup, renk_temasi, ts in gosterilecek_ozel_etkinlikler:
            t_etk = doc.add_table(rows=0, cols=2)
            t_etk.style = "Table Grid"
            add_merged_row(t_etk, f"{ts} | HALL: {ks}", renk_temasi, bold=True, align="center")
            for _, r in grup.iterrows():
                m = "\n".join(
                    [
                        str(r.get(c, "-")).strip()
                        for c in ["ana_baslik", "alt_baslik", "sol_metin", "sag_metin"]
                        if str(r.get(c, "-")).strip() not in ["-", "nan", ""]
                    ]
                )
                add_merged_row(t_etk, m, renk_temasi, align="center")
            doc.add_paragraph()

    t_bilimsel_baslik = doc.add_table(rows=0, cols=2)
    t_bilimsel_baslik.style = "Table Grid"
    add_merged_row(
        t_bilimsel_baslik,
        "--- BİLİMSEL BİLDİRİ PROGRAMI ---",
        "000000",
        text_color=(255, 255, 255),
        bold=True,
        align="center",
    )
    doc.add_paragraph()

    mevcut_islenen_gun = ""
    sirali_oturumlar = sorted(prog.items(), key=lambda x: parse_zaman(x[0][0]))

    for (oturum_zaman, sal), bilds in sirali_oturumlar:
        parcalar = oturum_zaman.split(" | ")
        t = parcalar[0] if len(parcalar) > 0 else oturum_zaman
        sa = parcalar[1] if len(parcalar) > 1 else "-"
        sn = bilds[0]["sid"] if bilds else "-"

        if t != mevcut_islenen_gun:
            mevcut_islenen_gun = t
            t_gun = doc.add_table(rows=0, cols=2)
            t_gun.style = "Table Grid"
            add_merged_row(t_gun, f">>> {t.upper()} BİLİMSEL PROGRAMI <<<", "B4C6E7", bold=True, align="center")
            add_merged_row(
                t_gun,
                "Her oturumdaki moderatör oturumu yönetecek ve Oturum Değerlendirici ile birbirinden bağımsız olarak EN İYİ BİLDİRİ (BEST PAPER) ÖDÜLLERi için bildiri sunumlarını değerlendireceklerdir.",
                "FFF2CC",
                align="center",
            )
            o_sayi = gun_istatistikleri.get(t, {}).get("oturum_sayisi", 0)
            b_sayi = gun_istatistikleri.get(t, {}).get("bildiri_sayisi", 0)
            add_merged_row(
                t_gun,
                f"Bugün toplam {o_sayi} adet bildiri sunum oturumu gerçekleşecek ve {b_sayi} adet bildiri sunulacaktır.",
                "E2EFDA",
                bold=True,
                text_color=(55, 86, 35),
                align="center",
            )
            doc.add_paragraph()

        mod_isim = mod_atamalari[(oturum_zaman, sal)]["mod"]
        deg_isim = mod_atamalari[(oturum_zaman, sal)]["deg"]
        oturum_metni = f"HALL: {sal}\n{sn}\n\n{bilds[0]['k'].upper()}\nModerator: {mod_isim}\nOturum Değerlendirici: {deg_isim}"

        t_oturum = doc.add_table(rows=0, cols=2)
        t_oturum.style = "Table Grid"
        add_merged_row(t_oturum, f"{t} | {sa}", "FFD966", bold=True)
        add_merged_row(t_oturum, oturum_metni, "FFE699", bold=True)

        for i, b in enumerate(bilds):
            bg_color = "F2F2F2" if i % 2 == 1 else "FFFFFF"
            row = t_oturum.add_row()
            yazarlar_metni = b["y"]
            sunucu = b["s"]
            cell_yazar = row.cells[0]
            cell_yazar.text = ""
            set_cell_bg(cell_yazar, bg_color)
            p = cell_yazar.paragraphs[0]
            if sunucu and sunucu not in ["-", "nan", ""] and sunucu in yazarlar_metni:
                parcalar_yazar = yazarlar_metni.split(sunucu, 1)
                if parcalar_yazar[0]:
                    p.add_run(parcalar_yazar[0])
                run_s = p.add_run(sunucu)
                run_s.underline = True
                run_s.bold = True
                if len(parcalar_yazar) > 1 and parcalar_yazar[1]:
                    p.add_run(parcalar_yazar[1])
            else:
                p.add_run(yazarlar_metni)
            row.cells[1].text = b["b"]
            set_cell_bg(row.cells[1], bg_color)
        doc.add_paragraph()

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def oturum_ozeti_olustur(df_ayar):
    df_ozet = df_ayar[["Gun_ve_Saat", "Salon", "Oturum_ID", "Sunum_Tipi"]].copy()
    df_ozet = df_ozet[~df_ozet["Gun_ve_Saat"].isin(["-", "ATANMADI", "İPTAL EDİLDİ", "nan", "NaN", ""])]
    df_ozet = df_ozet.dropna(subset=["Gun_ve_Saat"])
    df_ozet["Salon"] = df_ozet["Salon"].apply(clean_salon)
    df_grouped = (
        df_ozet.groupby(["Gun_ve_Saat", "Salon", "Oturum_ID", "Sunum_Tipi"])
        .size()
        .reset_index(name="Atanan_Bildiri_Sayisi")
    )
    df_grouped["sort_time"] = df_grouped["Gun_ve_Saat"].apply(parse_zaman)
    df_grouped = df_grouped.sort_values(by=["sort_time", "Salon"]).drop(columns=["sort_time"])

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df_grouped.to_excel(writer, index=False, sheet_name="Oturum_Rontgeni")
        workbook = writer.book
        worksheet = writer.sheets["Oturum_Rontgeni"]
        baslik_formati = workbook.add_format({"bold": True, "bg_color": "#4472C4", "font_color": "white", "border": 1})
        for col_num, value in enumerate(df_grouped.columns.values):
            worksheet.write(0, col_num, value, baslik_formati)
        worksheet.set_column("A:A", 35)
        worksheet.set_column("B:B", 25)
        worksheet.set_column("C:C", 15)
        worksheet.set_column("D:D", 15)
        worksheet.set_column("E:E", 20)
    output.seek(0)
    return output


def excel_bas(df_bildiriler, df_ozel, df_mod, tip):
    prog = build_program_dict(df_bildiriler)
    tum_zamanlar = set()
    tum_salonlar = set()
    for kz, ks in prog.keys():
        tum_zamanlar.add(kz)
        tum_salonlar.add(ks)

    gosterilecek_ozel_etkinlikler = []
    soft_palette = ["#DDEBF7", "#FCE4D6", "#E2EFDA", "#FFF2CC", "#E6E6FA", "#F2F2F2"]
    if df_ozel is not None and not df_ozel.empty and {"oturum_sirasi", "salon"}.issubset(set(df_ozel.columns)):
        color_index = 0
        for (sira, sal), grup in df_ozel.groupby(["oturum_sirasi", "salon"], sort=False):
            ks = clean_salon(sal)
            tum_salonlar.add(ks)
            renk_temasi = soft_palette[color_index % len(soft_palette)]
            color_index += 1
            ts = str(grup.iloc[0].get("tarih_saat", "-")).strip().replace("Persembe", "Perşembe")
            t_event = "07.05.2026 Perşembe" if "07" in ts else ("08.05.2026 Cuma" if "08" in ts else "Bilinmeyen Gun")
            m = re.search(r"(\d{2}[.:]\d{2})\s*-\s*(\d{2}[.:]\d{2})", ts)
            sa_event = f"{m.group(1).replace('.', ':')}-{m.group(2).replace('.', ':')}" if m else ts
            zaman_key = f"{t_event} | {sa_event}"
            tum_zamanlar.add(zaman_key)
            gosterilecek_ozel_etkinlikler.append((sira, ks, grup, renk_temasi, zaman_key))

    mod_atamalari = moderator_atamalari_olustur(prog, df_mod, tip)
    gun_istatistikleri = gun_istatistikleri_olustur(prog)

    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output, {"in_memory": True})

    map_data = {}
    for (kz, ks), bilds in prog.items():
        if (kz, ks) not in map_data:
            map_data[(kz, ks)] = {"session_count": 0, "event_name": None, "event_color": None}
        map_data[(kz, ks)]["session_count"] = len(bilds)

    for _sira, ks, grup, renk, zaman_key in gosterilecek_ozel_etkinlikler:
        ana_baslik = str(grup.iloc[0].get("ana_baslik", "Etkinlik")).strip()
        if (zaman_key, ks) not in map_data:
            map_data[(zaman_key, ks)] = {"session_count": 0, "event_name": None, "event_color": None}
        map_data[(zaman_key, ks)]["event_name"] = ana_baslik
        map_data[(zaman_key, ks)]["event_color"] = renk

    list_zamanlar = sorted(list(tum_zamanlar), key=parse_zaman)
    list_salonlar = sorted(list(tum_salonlar))

    ws_harita = workbook.add_worksheet("Oturum_Haritasi")
    f_map_baslik = workbook.add_format({"bold": True, "bg_color": "#4472C4", "font_color": "white", "border": 1, "align": "center", "valign": "vcenter"})
    f_map_saat = workbook.add_format({"bold": True, "bg_color": "#D9E1F2", "border": 1, "align": "left"})
    f_map_bos = workbook.add_format({"bg_color": "#F2F2F2", "font_color": "#A6A6A6", "border": 1, "align": "center", "valign": "vcenter"})
    f_map_dolu = workbook.add_format({"bold": True, "bg_color": "#E2EFDA", "font_color": "#375623", "border": 1, "align": "center", "valign": "vcenter"})
    f_map_cakisma = workbook.add_format({"bold": True, "bg_color": "#FF0000", "font_color": "white", "border": 1, "align": "center", "valign": "vcenter", "text_wrap": True})
    dinamik_renk_cache = {}

    ws_harita.set_column(0, 0, 35)
    ws_harita.write(0, 0, "SAAT / OTURUM", f_map_baslik)
    for col_idx, sal in enumerate(list_salonlar, 1):
        ws_harita.set_column(col_idx, col_idx, 25)
        ws_harita.write(0, col_idx, sal, f_map_baslik)

    for row_idx, z_key in enumerate(list_zamanlar, 1):
        ws_harita.write(row_idx, 0, z_key, f_map_saat)
        for col_idx, sal in enumerate(list_salonlar, 1):
            hucre = map_data.get((z_key, sal), {"session_count": 0, "event_name": None})
            s_count = hucre["session_count"]
            e_name = hucre["event_name"]
            e_color = hucre.get("event_color", "#FFD966")

            if s_count > 0 and e_name:
                ws_harita.write(row_idx, col_idx, f"ÇAKIŞMA!\n[{s_count} Bildiri] VE [{e_name}]", f_map_cakisma)
                ws_harita.set_row(row_idx, 45)
            elif s_count > 0:
                ws_harita.write(row_idx, col_idx, f"{s_count} Bildiri Var", f_map_dolu)
            elif e_name:
                if e_color not in dinamik_renk_cache:
                    dinamik_renk_cache[e_color] = workbook.add_format({"bold": True, "bg_color": e_color, "border": 1, "align": "center", "valign": "vcenter", "text_wrap": True})
                ws_harita.write(row_idx, col_idx, f"[ETKİNLİK]\n{e_name}", dinamik_renk_cache[e_color])
                ws_harita.set_row(row_idx, 40)
            else:
                ws_harita.write(row_idx, col_idx, "BOŞ", f_map_bos)

    worksheet = workbook.add_worksheet("Kesin_Program")
    worksheet.set_column("A:A", 65)
    worksheet.set_column("B:B", 95)

    fmt = {
        "ayirici": workbook.add_format({"bg_color": "#000000", "font_color": "#FFFFFF", "bold": True, "align": "center", "valign": "vcenter"}),
        "baslik_dev": workbook.add_format({"bold": True, "font_size": 14, "align": "center", "valign": "vcenter", "border": 1}),
        "mavi_gun": workbook.add_format({"bold": True, "bg_color": "#B4C6E7", "align": "left", "valign": "vcenter", "border": 1, "text_wrap": True}),
        "odul_sari": workbook.add_format({"bg_color": "#FFF2CC", "align": "left", "valign": "vcenter", "border": 1, "text_wrap": True}),
        "istatistik_yesil": workbook.add_format({"bold": True, "bg_color": "#E2EFDA", "font_color": "#375623", "align": "center", "valign": "vcenter", "border": 1, "text_wrap": True}),
        "saat_baslik": workbook.add_format({"bold": True, "bg_color": "#FFD966", "align": "left", "valign": "vcenter", "border": 1, "text_wrap": True}),
        "salon_baslik": workbook.add_format({"bold": True, "bg_color": "#FFE699", "align": "left", "valign": "vcenter", "border": 1, "text_wrap": True}),
        "bildiri_tek": workbook.add_format({"bg_color": "#FFFFFF", "align": "left", "valign": "vcenter", "border": 1, "text_wrap": True}),
        "bildiri_tek_u": workbook.add_format({"bold": True, "underline": True, "bg_color": "#FFFFFF", "border": 1, "text_wrap": True}),
        "bildiri_tek_u_only": workbook.add_format({"bold": True, "underline": True}),
        "bildiri_cift": workbook.add_format({"bg_color": "#F2F2F2", "align": "left", "valign": "vcenter", "border": 1, "text_wrap": True}),
        "bildiri_cift_u": workbook.add_format({"bold": True, "underline": True, "bg_color": "#F2F2F2", "border": 1, "text_wrap": True}),
        "bildiri_cift_u_only": workbook.add_format({"bold": True, "underline": True}),
    }

    satir_no = 0
    ana_baslik = "11. IHMC FİZİKİ / FACE TO FACE PROGRAM" if tip == "YUZYUZE" else "11. IHMC ONLINE (DIGITAL) PROGRAM"
    worksheet.merge_range(satir_no, 0, satir_no, 1, ana_baslik, fmt["baslik_dev"])
    worksheet.set_row(satir_no, 40)
    satir_no += 2

    if len(gosterilecek_ozel_etkinlikler) > 0:
        worksheet.merge_range(satir_no, 0, satir_no, 1, "ÖZEL ETKİNLİKLER", fmt["ayirici"])
        satir_no += 1
        for _sira, ks, grup, renk_temasi, _z_key in gosterilecek_ozel_etkinlikler:
            ts = str(grup.iloc[0].get("tarih_saat", "-")).strip().replace("07.05.2026", "07.05.2026 Perşembe /").replace("08.05.2026", "08.05.2026 Cuma /")
            if renk_temasi not in dinamik_renk_cache:
                dinamik_renk_cache[renk_temasi] = workbook.add_format({"bold": True, "bg_color": renk_temasi, "align": "left", "valign": "vcenter", "border": 1, "text_wrap": True})
            dinamik_f = dinamik_renk_cache[renk_temasi]
            worksheet.merge_range(satir_no, 0, satir_no, 1, f"{ts} | HALL: {ks}", dinamik_f)
            satir_no += 1
            for _, r in grup.iterrows():
                metin = "\n".join(
                    [
                        str(r.get(c, "-")).strip()
                        for c in ["ana_baslik", "alt_baslik", "sol_metin", "sag_metin"]
                        if str(r.get(c, "-")).strip() not in ["-", "nan", ""]
                    ]
                )
                worksheet.merge_range(satir_no, 0, satir_no, 1, metin, dinamik_f)
                worksheet.set_row(satir_no, 50)
                satir_no += 1
            satir_no += 1

    worksheet.merge_range(satir_no, 0, satir_no, 1, "--- BİLİMSEL BİLDİRİ PROGRAMI ---", fmt["ayirici"])
    satir_no += 2

    mevcut_islenen_gun = ""
    sirali_oturumlar = sorted(prog.items(), key=lambda x: parse_zaman(x[0][0]))

    for (oturum_zaman, sal), bilds in sirali_oturumlar:
        parcalar = oturum_zaman.split(" | ")
        t = parcalar[0] if len(parcalar) > 0 else oturum_zaman
        sa = parcalar[1] if len(parcalar) > 1 else "-"
        sn = bilds[0]["sid"] if bilds else "-"

        if t != mevcut_islenen_gun:
            mevcut_islenen_gun = t
            worksheet.merge_range(satir_no, 0, satir_no, 1, f">>> {t.upper()} BİLİMSEL PROGRAMI <<<", fmt["mavi_gun"])
            worksheet.set_row(satir_no, 40)
            satir_no += 1

            odul_notu = "Her oturumdaki moderatör oturumu yönetecek ve Oturum Değerlendirici ile birbirinden bağımsız olarak EN İYİ BİLDİRİ (BEST PAPER) ÖDÜLLERi için bildiri sunumlarını değerlendireceklerdir."
            worksheet.merge_range(satir_no, 0, satir_no, 1, odul_notu, fmt["odul_sari"])
            worksheet.set_row(satir_no, 45)
            satir_no += 1

            o_sayi = gun_istatistikleri.get(t, {}).get("oturum_sayisi", 0)
            b_sayi = gun_istatistikleri.get(t, {}).get("bildiri_sayisi", 0)
            istatistik_notu = f"Bugün toplam {o_sayi} adet bildiri sunum oturumu gerçekleşecek ve {b_sayi} adet bildiri sunulacaktır."
            worksheet.merge_range(satir_no, 0, satir_no, 1, istatistik_notu, fmt["istatistik_yesil"])
            worksheet.set_row(satir_no, 25)
            satir_no += 2

        mod_isim = mod_atamalari[(oturum_zaman, sal)]["mod"]
        deg_isim = mod_atamalari[(oturum_zaman, sal)]["deg"]
        oturum_metni = f"HALL: {sal}\n{sn}\n\n{bilds[0]['k'].upper()}\nModerator: {mod_isim}\nOturum Değerlendirici: {deg_isim}"

        worksheet.merge_range(satir_no, 0, satir_no, 1, f"{t} | {sa}", fmt["saat_baslik"])
        satir_no += 1
        worksheet.merge_range(satir_no, 0, satir_no, 1, oturum_metni, fmt["salon_baslik"])
        worksheet.set_row(satir_no, 100)
        satir_no += 1

        for i, b in enumerate(bilds):
            yazarlar_metni, sunucu = b["y"], b["s"]
            is_cift = i % 2 == 1
            f_norm = fmt["bildiri_cift"] if is_cift else fmt["bildiri_tek"]
            f_ul_cell = fmt["bildiri_cift_u"] if is_cift else fmt["bildiri_tek_u"]
            f_ul_txt = fmt["bildiri_cift_u_only"] if is_cift else fmt["bildiri_tek_u_only"]

            if sunucu == yazarlar_metni and sunucu not in ["-", "nan", ""]:
                worksheet.write(satir_no, 0, yazarlar_metni, f_ul_cell)
            elif sunucu in yazarlar_metni and sunucu not in ["-", "nan", ""]:
                parcalar_yazar = yazarlar_metni.split(sunucu, 1)
                rich_text = []
                if parcalar_yazar[0]:
                    rich_text.extend([f_norm, parcalar_yazar[0]])
                rich_text.extend([f_ul_txt, sunucu])
                if len(parcalar_yazar) > 1 and parcalar_yazar[1]:
                    rich_text.extend([f_norm, parcalar_yazar[1]])
                if len(rich_text) >= 2:
                    worksheet.write_rich_string(satir_no, 0, *rich_text, f_norm)
                else:
                    worksheet.write(satir_no, 0, yazarlar_metni, f_norm)
            else:
                worksheet.write(satir_no, 0, yazarlar_metni, f_norm)

            worksheet.write(satir_no, 1, b["b"], f_norm)
            satir_no += 1
        satir_no += 1

    workbook.close()
    output.seek(0)
    return output


# =========================
# RAPOR MOTORLARI
# =========================

def make_genel_report(katilimcilar):
    output = io.BytesIO()
    writer = pd.ExcelWriter(output, engine="xlsxwriter")
    rapor_list = []
    for _, v in katilimcilar.items():
        odeme = turkce_buyuk(v["Odeme"])
        gorev = turkce_buyuk(v["Gorev"])
        b_list = v["Bildiriler"]
        renk = "WHITE"
        if "MUAF" in odeme or "GÖREVLİ" in odeme or "DAVETLİ" in gorev:
            renk = "PURPLE"
        elif "İNDİRİM" in odeme or "ÖĞRENCİ" in odeme:
            renk = "ORANGE"
        elif "BEKLENİYOR" in odeme and len(b_list) > 0:
            renk = "YELLOW"
        elif "BEKLENİYOR" not in odeme and "YOK" not in odeme and len(b_list) > 0:
            renk = "GREEN"
        elif "BEKLENİYOR" not in odeme and "YOK" not in odeme and len(b_list) == 0:
            renk = "BLUE"
        rapor_list.append(
            {
                "Unvan": turkce_buyuk(v["Unvan"]),
                "Adı Soyadı": v["Orijinal İsim"],
                "Ödeme Durumu": odeme,
                "Bildiri 1": turkce_buyuk(b_list[0]) if len(b_list) > 0 else "",
                "Bildiri 2": turkce_buyuk(b_list[1]) if len(b_list) > 1 else "",
                "Katılım Türü": turkce_buyuk(v["Tur"]),
                "_Color": renk,
            }
        )
    df_r = pd.DataFrame(rapor_list).sort_values(by="Adı Soyadı") if rapor_list else pd.DataFrame()
    if df_r.empty:
        pd.DataFrame(columns=["Unvan", "Adı Soyadı", "Ödeme Durumu", "Bildiri 1", "Bildiri 2", "Katılım Türü"]).to_excel(writer, sheet_name="Genel Katılımcı Listesi", index=False)
    else:
        df_r.drop(columns=["_Color"]).to_excel(writer, sheet_name="Genel Katılımcı Listesi", index=False)
        workbook = writer.book
        ws = writer.sheets["Genel Katılımcı Listesi"]
        fmts = {
            "GREEN": workbook.add_format({"bg_color": "#C6EFCE"}),
            "BLUE": workbook.add_format({"bg_color": "#BDD7EE"}),
            "YELLOW": workbook.add_format({"bg_color": "#FFEB9C"}),
            "PURPLE": workbook.add_format({"bg_color": "#E4DFEC"}),
            "ORANGE": workbook.add_format({"bg_color": "#FCE4D6"}),
        }
        ws.set_column("B:B", 30)
        ws.set_column("C:E", 40)
        for i, r_code in enumerate(df_r["_Color"]):
            if r_code in fmts:
                ws.set_row(i + 1, cell_format=fmts[r_code])
    writer.close()
    return output.getvalue()


def make_unpaid_report(bildiriler, katilimcilar):
    bekleyen = []
    for _, bv in bildiriler.items():
        if bv["Odeme"] == "Hayır":
            yazarlar_str = ", ".join([turkce_buyuk(katilimcilar.get(y, {}).get("Orijinal İsim", y)) for y in bv["Yazarlar"]])
            bekleyen.append({"KABUL EDİLMİŞ AMA ÖDENMEMİŞ BİLDİRİ": turkce_buyuk(bv["Orijinal İsim"]), "YAZARLARI": yazarlar_str})
    out = io.BytesIO()
    pd.DataFrame(bekleyen).to_excel(out, index=False)
    return out.getvalue()


def make_katilim_turu_report(katilimcilar, tur_tipi):
    liste = []
    for _, v in katilimcilar.items():
        if super_temiz(tur_tipi) in super_temiz(v["Tur"]):
            liste.append(
                {
                    "AD SOYAD": v["Orijinal İsim"],
                    "UNVAN": turkce_buyuk(v["Unvan"]),
                    "KATILIM TÜRÜ": turkce_buyuk(v["Tur"]),
                    "ÖDEME DURUMU": v["Odeme"],
                    "TELEFON": v["Telefon"],
                    "E-POSTA": v["E-Posta"],
                }
            )
    out = io.BytesIO()
    if liste:
        pd.DataFrame(liste).sort_values(by="AD SOYAD").to_excel(out, index=False)
    else:
        pd.DataFrame(columns=["AD SOYAD", "UNVAN", "KATILIM TÜRÜ", "ÖDEME DURUMU", "TELEFON", "E-POSTA"]).to_excel(out, index=False)
    return out.getvalue()


def make_meal_report(katilimcilar, keys):
    meals = []
    for _, v in katilimcilar.items():
        etk = super_temiz(v["Etkinlikler"])
        if all(super_temiz(x) in etk for x in keys):
            meals.append({"AD SOYAD": v["Orijinal İsim"], "TUR": turkce_buyuk(v["Tur"]), "ODEME": turkce_buyuk(v["Odeme"])})
    out = io.BytesIO()
    if meals:
        pd.DataFrame(meals).to_excel(out, index=False)
    else:
        pd.DataFrame(columns=["AD SOYAD", "TUR", "ODEME"]).to_excel(out, index=False)
    return out.getvalue()


# =========================
# STREAMLIT UI
# =========================

st.set_page_config(page_title="IHMC 2026 Süper Kongre Paneli", layout="wide", page_icon="📊")

st.markdown(
    """
    <style>
    .main { background-color: #f8fafc; }
    .stButton>button, .stDownloadButton>button { width: 100%; border-radius: 6px; font-weight: 700; }
    .metric-card { padding: 1rem; border-radius: 8px; background: white; border: 1px solid #e2e8f0; }
    div[data-testid="stFormSubmitButton"] button {
        background-color: #2563eb !important;
        color: white !important;
        border: none !important;
        height: 3em;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


def sidebar_settings():
    st.sidebar.title("Ayarlar")
    kat_id = st.sidebar.text_input("Katılımcı veri Sheet ID", value=DEFAULT_KATILIMCI_SHEET_ID)
    prog_id = st.sidebar.text_input("Program/Matbaa Sheet ID", value=DEFAULT_PROGRAM_SHEET_ID)
    writable = has_service_account() and get_gspread_client() is not None
    if writable:
        st.sidebar.success("Google Sheets yazma modu aktif")
    else:
        st.sidebar.warning("Sadece okuma modu. Yazmak için gcp_service_account secrets gerekli.")
    if st.sidebar.button("Verileri Yenile"):
        st.cache_data.clear()
        st.rerun()
    st.sidebar.caption("Yazma için iki Google Sheet dosyasını servis hesabı e-postasıyla paylaşın.")
    return kat_id, prog_id, writable


def row_label(row, cols):
    parts = [f"Satır {row.get('_row', '?')}"]
    for col in cols:
        if col in row and safe_str(row[col]):
            parts.append(safe_str(row[col])[:90])
    return " | ".join(parts)


katilimci_sheet_id, program_sheet_id, writable = sidebar_settings()

st.title("IHMC 2026 Süper Kongre Paneli")
st.caption("Program yönetimi, matbaa çıktıları, katılımcı sorgulama, moderatör yönetimi ve raporlar tek ekranda.")

try:
    with st.spinner("Google Sheets verileri yükleniyor..."):
        raw_data = load_all_data(katilimci_sheet_id, program_sheet_id, writable)
    bildiriler, katilimcilar = build_database(raw_data)
    df_master = build_matbaa_master(raw_data)
except Exception as e:
    st.error(f"Veri bağlantı hatası: {e}")
    st.stop()


tab_search, tab_program, tab_salon, tab_mod, tab_matbaa, tab_reports, tab_telegram = st.tabs(
    [
        "Hızlı Sorgu",
        "Bildiri Programı",
        "Salon Kontrolü",
        "Moderatörler",
        "Matbaa",
        "Raporlar",
        "Telegram",
    ]
)


with tab_search:
    st.subheader("Katılımcı veya Bildiri Ara")
    if "arama_metni" not in st.session_state:
        st.session_state.arama_metni = ""
    with st.form("search_form"):
        c1, c2 = st.columns([4, 1])
        sorgu_input = c1.text_input("Arama", value=st.session_state.arama_metni, placeholder="Örn: Cihan veya Yapay Zeka", label_visibility="collapsed")
        arama = c2.form_submit_button("Ara")
    if arama:
        st.session_state.arama_metni = sorgu_input

    sorgu_aktif = st.session_state.arama_metni
    if sorgu_aktif:
        t_sorgu = temiz_metin(sorgu_aktif)
        found_kisiler = [k for k in katilimcilar.keys() if t_sorgu in k]
        found_bildiriler = [b for b in bildiriler.keys() if t_sorgu in b]

        if found_kisiler:
            st.write(f"### Kişi Sonuçları ({len(found_kisiler)})")
            for k_key in found_kisiler:
                kv = katilimcilar[k_key]
                with st.expander(f"{get_kisi_renk_emoji(kv)} {turkce_buyuk(kv['Unvan'])} {kv['Orijinal İsim']}", expanded=True):
                    c1, c2 = st.columns(2)
                    c1.write(f"**Ödeme:** {kv['Odeme']}")
                    c1.write(f"**Görev:** {kv['Gorev']}")
                    c1.write(f"**Katılım Türü:** {kv['Tur']}")
                    c1.write(f"**Telefon:** {kv['Telefon'] or 'Kayıtlı değil'}")
                    c1.write(f"**E-posta:** {kv['E-Posta'] or 'Kayıtlı değil'}")
                    yemekler = ", ".join([y.strip() for y in super_temiz(kv["Etkinlikler"]).split(",") if y.strip()])
                    c2.write(f"**Etkinlikler:** {yemekler if yemekler else 'Yok'}")
                    c2.write(f"**Günler:** 6 Mayıs: {kv['Gunler']['6']} | 7 Mayıs: {kv['Gunler']['7']} | 8 Mayıs: {kv['Gunler']['8']}")
                    st.write("**Bildirileri ve Programı:**")
                    if kv["Bildiriler"]:
                        for b in kv["Bildiriler"]:
                            b_veri = bildiriler.get(temiz_metin(b), {})
                            st.markdown(f"- {get_bildiri_renk_emoji(b_veri)} **{b}**")
                            st.caption(get_program_info(b, raw_data.get("program_plan", pd.DataFrame())))
                    else:
                        st.write("Bildirisi yok.")

        if found_bildiriler:
            st.write(f"### Bildiri Sonuçları ({len(found_bildiriler)})")
            for b_key in found_bildiriler:
                bv = bildiriler[b_key]
                with st.expander(f"{get_bildiri_renk_emoji(bv)} {bv['Orijinal İsim']}", expanded=True):
                    st.write(f"**Konu:** {bv['Konu']} | **Butik:** {bv['Butik']}")
                    st.write(f"**Ödeme:** {'Evet, Yapıldı' if bv.get('Odeme') == 'Evet' else 'Hayır, Bekliyor'}")
                    st.write(f"**Program:** {get_program_info(bv['Orijinal İsim'], raw_data.get('program_plan', pd.DataFrame()))}")
                    st.write("**Yazarlar:**")
                    for y in bv["Yazarlar"]:
                        y_kisi = katilimcilar.get(y, {})
                        st.write(f"{get_kisi_renk_emoji(y_kisi)} {turkce_buyuk(y_kisi.get('Unvan', ''))} {turkce_buyuk(y_kisi.get('Orijinal İsim', y))}")

        if not found_kisiler and not found_bildiriler:
            st.warning("Sonuç bulunamadı.")


with tab_program:
    st.subheader("Bildiri Gün/Saat/Salon/Oturum/Sunum Tipi Yönetimi")
    df_plan = raw_data.get("program_plan", pd.DataFrame()).copy()
    df_bild = raw_data.get("program_bildiriler", pd.DataFrame()).copy()

    required_cols = ["Bildiri_Adi", "Gun_ve_Saat", "Salon", "Oturum_ID", "Sunum_Tipi"]
    missing_required = [c for c in required_cols if c not in df_plan.columns]
    if missing_required:
        st.error(f"Program planı sheet'inde eksik kolonlar var: {', '.join(missing_required)}")
    else:
        f1, f2, f3 = st.columns(3)
        q = f1.text_input("Bildiri adı filtrele", "")
        tip_filter = f2.multiselect("Sunum tipi", sorted([x for x in df_plan["Sunum_Tipi"].dropna().unique() if safe_str(x)]))
        gunler = sorted({gun_key(x) for x in df_plan["Gun_ve_Saat"].dropna().astype(str) if safe_str(x)}, key=parse_zaman)
        gun_filter = f3.multiselect("Gün", gunler)

        df_show = df_plan.copy()
        if q:
            df_show = df_show[df_show["Bildiri_Adi"].astype(str).apply(lambda x: temiz_metin(q) in temiz_metin(x))]
        if tip_filter:
            df_show = df_show[df_show["Sunum_Tipi"].isin(tip_filter)]
        if gun_filter:
            df_show = df_show[df_show["Gun_ve_Saat"].astype(str).apply(lambda x: gun_key(x) in gun_filter)]

        st.dataframe(
            df_show[[c for c in ["_row", "Bildiri_Adi", "Sunum_Tipi", "Gun_ve_Saat", "Salon", "Oturum_ID"] if c in df_show.columns]],
            use_container_width=True,
            hide_index=True,
        )

        if not df_show.empty:
            options = {row_label(r, ["Bildiri_Adi", "Sunum_Tipi", "Gun_ve_Saat", "Salon", "Oturum_ID"]): r for _, r in df_show.iterrows()}
            selected_label = st.selectbox("Düzenlenecek bildiriyi seç", list(options.keys()))
            selected = options[selected_label]
            selected_bildiri = safe_str(selected["Bildiri_Adi"])

            existing_sessions = (
                df_plan[["Gun_ve_Saat", "Salon", "Oturum_ID", "Sunum_Tipi"]]
                .drop_duplicates()
                .sort_values(by=["Gun_ve_Saat", "Salon"], key=lambda s: s.astype(str))
            )
            session_labels = [
                f"{r['Sunum_Tipi']} | {r['Gun_ve_Saat']} | {r['Salon']} | {r['Oturum_ID']}"
                for _, r in existing_sessions.iterrows()
            ]

            st.write("#### Seçili Bildiriyi Taşı / Güncelle")
            with st.form("program_update_form"):
                use_existing = st.checkbox("Mevcut bir oturuma taşı", value=True)
                if use_existing and session_labels:
                    current_label = f"{selected['Sunum_Tipi']} | {selected['Gun_ve_Saat']} | {selected['Salon']} | {selected['Oturum_ID']}"
                    default_index = session_labels.index(current_label) if current_label in session_labels else 0
                    chosen_session = st.selectbox("Hedef oturum", session_labels, index=default_index)
                    chosen_parts = chosen_session.split(" | ", 3)
                    new_tip, new_time, new_salon, new_oturum = chosen_parts[0], chosen_parts[1], chosen_parts[2], chosen_parts[3]
                else:
                    c1, c2 = st.columns(2)
                    new_tip = c1.selectbox("Sunum tipi", ["Yuzyuze", "Online"], index=0 if safe_str(selected["Sunum_Tipi"]) != "Online" else 1)
                    new_time = c2.text_input("Gün ve saat", value=safe_str(selected["Gun_ve_Saat"]))
                    c3, c4 = st.columns(2)
                    salon_options = sorted({safe_str(x) for x in df_plan["Salon"].dropna().unique() if safe_str(x)})
                    if safe_str(selected["Salon"]) not in salon_options:
                        salon_options.insert(0, safe_str(selected["Salon"]))
                    new_salon = c3.selectbox("Salon", salon_options) if salon_options else c3.text_input("Salon", value=safe_str(selected["Salon"]))
                    new_oturum = c4.text_input("Oturum ID", value=safe_str(selected["Oturum_ID"]))

                current_authors = bildiriler.get(temiz_metin(selected_bildiri), {}).get("Yazarlar", [])
                current_presenter = bildiriler.get(temiz_metin(selected_bildiri), {}).get("Sunucu", "")
                author_display = [turkce_buyuk(a) for a in current_authors]
                author_map = {turkce_buyuk(a): a for a in current_authors}
                presenter_choice = st.selectbox(
                    "Sunacak yazar",
                    ["Değiştirme"] + author_display + ["Elle yaz"],
                    index=0,
                )
                presenter_free = ""
                if presenter_choice == "Elle yaz":
                    presenter_free = st.text_input("Yeni sunucu adı", value=current_presenter)

                kaydet = st.form_submit_button("Bildiri Güncelle")

            if kaydet:
                if not writable:
                    st.error("Yazma modu aktif değil. Önce Streamlit secrets içine servis hesabı ekleyin.")
                else:
                    try:
                        update_sheet_row(
                            program_sheet_id,
                            PROGRAM_PLAN_INDEX,
                            int(selected["_row"]),
                            {
                                "Gun_ve_Saat": new_time,
                                "Salon": new_salon,
                                "Oturum_ID": new_oturum,
                                "Sunum_Tipi": new_tip,
                            },
                        )
                        if presenter_choice != "Değiştirme":
                            new_presenter = presenter_free if presenter_choice == "Elle yaz" else author_map[presenter_choice]
                            b_name_col = find_col(df_bild, ["Bildiri Ismi", "Bildiri_Adi", "Bildiri Adı"], contains=["bildiri"])
                            presenter_col = find_col(df_bild, ["Sunan Yazar", "sunucu", "Sunucu"], contains=["sunan"])
                            if b_name_col and presenter_col:
                                b_row = find_row_by_value(program_sheet_id, PROGRAM_BILDIRILER_SHEET, b_name_col, selected_bildiri)
                                if b_row:
                                    update_sheet_row(program_sheet_id, PROGRAM_BILDIRILER_SHEET, b_row, {presenter_col: new_presenter})
                        st.cache_data.clear()
                        st.success("Bildiri güncellendi.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Güncelleme hatası: {e}")


with tab_salon:
    st.subheader("Ekranda Salon ve Oturum Kapasite Kontrolü")
    df_plan = raw_data.get("program_plan", pd.DataFrame()).copy()
    if df_plan.empty or "Gun_ve_Saat" not in df_plan.columns or "Salon" not in df_plan.columns:
        st.warning("Program planı bulunamadı.")
    else:
        valid = df_plan[~df_plan["Gun_ve_Saat"].isin(["", "-", "nan", "NaN", "ATANMADI", "İPTAL EDİLDİ"])].copy()
        valid["Gün"] = valid["Gun_ve_Saat"].apply(gun_key)
        gunler = sorted(valid["Gün"].dropna().unique(), key=parse_zaman)
        secili_gun = st.selectbox("Kontrol edilecek gün", ["Tümü"] + list(gunler))
        if secili_gun != "Tümü":
            valid = valid[valid["Gün"] == secili_gun]

        c1, c2, c3 = st.columns(3)
        c1.metric("Toplam Bildiri", len(valid))
        c2.metric("Aktif Salon", valid["Salon"].nunique())
        c3.metric("Oturum", valid[["Gun_ve_Saat", "Salon", "Oturum_ID"]].drop_duplicates().shape[0] if "Oturum_ID" in valid.columns else valid[["Gun_ve_Saat", "Salon"]].drop_duplicates().shape[0])

        counts = valid.groupby("Salon").size().reset_index(name="Bildiri Sayısı").sort_values("Bildiri Sayısı", ascending=False)
        st.write("#### Salonlara Göre Bildiri Sayısı")
        st.dataframe(counts, use_container_width=True, hide_index=True)
        st.bar_chart(counts.set_index("Salon"))

        st.write("#### Saat x Salon Haritası")
        pivot = valid.pivot_table(index="Gun_ve_Saat", columns="Salon", values="Bildiri_Adi", aggfunc="count", fill_value=0)
        pivot = pivot.sort_index(key=lambda idx: idx.map(parse_zaman))
        st.dataframe(pivot, use_container_width=True)

        st.write("#### Oturum Detayı")
        detail_cols = [c for c in ["Gun_ve_Saat", "Salon", "Oturum_ID", "Sunum_Tipi", "Bildiri_Adi"] if c in valid.columns]
        st.dataframe(valid[detail_cols].sort_values(by=["Gun_ve_Saat", "Salon"], key=lambda s: s.astype(str)), use_container_width=True, hide_index=True)


with tab_mod:
    st.subheader("Moderatör ve Oturum Değerlendirici Yönetimi")
    df_mod = raw_data.get("mod", pd.DataFrame()).copy()
    df_plan = raw_data.get("program_plan", pd.DataFrame()).copy()
    if df_mod.empty:
        st.warning("Moderatorler sheet'i bulunamadı veya boş.")
    else:
        view_cols = [c for c in ["_row", "unvan_ad_soyad", "kurum", "Gorev", "Gun_ve_Saat", "Salon", "Oturum_ID", "Online"] if c in df_mod.columns]
        st.dataframe(df_mod[view_cols], use_container_width=True, hide_index=True)

        mod_tab1, mod_tab2 = st.tabs(["Tek Kişi Değiştir", "İki Kişiyi Yer Değiştir"])

        with mod_tab1:
            options = {row_label(r, ["unvan_ad_soyad", "Gorev", "Gun_ve_Saat", "Salon", "Oturum_ID"]): r for _, r in df_mod.iterrows()}
            selected_label = st.selectbox("Düzenlenecek moderatör/değerlendirici", list(options.keys()))
            selected = options[selected_label]

            session_labels = []
            if not df_plan.empty and {"Gun_ve_Saat", "Salon", "Oturum_ID", "Sunum_Tipi"}.issubset(df_plan.columns):
                sessions = df_plan[["Gun_ve_Saat", "Salon", "Oturum_ID", "Sunum_Tipi"]].drop_duplicates()
                session_labels = [
                    f"{r['Sunum_Tipi']} | {r['Gun_ve_Saat']} | {r['Salon']} | {r['Oturum_ID']}"
                    for _, r in sessions.iterrows()
                ]

            with st.form("mod_update_form"):
                c1, c2 = st.columns(2)
                new_name = c1.text_input("Ad Soyad / Ünvan", value=safe_str(selected.get("unvan_ad_soyad", "")))
                new_kurum = c2.text_input("Kurum", value=safe_str(selected.get("kurum", "")))
                c3, c4 = st.columns(2)
                new_gorev = c3.text_input("Görev", value=safe_str(selected.get("Gorev", "")))
                new_online = c4.selectbox("Online", ["Hayır", "Evet"], index=1 if safe_str(selected.get("Online", "")).lower() in ["evet", "online", "true", "1"] else 0)

                use_existing_session = st.checkbox("Mevcut oturuma ata", value=True)
                if use_existing_session and session_labels:
                    current_session = f"{'Online' if new_online == 'Evet' else 'Yuzyuze'} | {safe_str(selected.get('Gun_ve_Saat', ''))} | {safe_str(selected.get('Salon', ''))} | {safe_str(selected.get('Oturum_ID', ''))}"
                    idx = session_labels.index(current_session) if current_session in session_labels else 0
                    chosen_session = st.selectbox("Hedef oturum", session_labels, index=idx)
                    p = chosen_session.split(" | ", 3)
                    _tip, new_time, new_salon, new_oturum = p[0], p[1], p[2], p[3]
                    new_online = "Evet" if _tip == "Online" else "Hayır"
                else:
                    c5, c6 = st.columns(2)
                    new_time = c5.text_input("Gün ve saat", value=safe_str(selected.get("Gun_ve_Saat", "")))
                    new_salon = c6.text_input("Salon", value=safe_str(selected.get("Salon", "")))
                    new_oturum = st.text_input("Oturum ID", value=safe_str(selected.get("Oturum_ID", "")))

                mod_kaydet = st.form_submit_button("Moderatör Kaydet")

            if mod_kaydet:
                if not writable:
                    st.error("Yazma modu aktif değil.")
                else:
                    try:
                        update_sheet_row(
                            program_sheet_id,
                            PROGRAM_MOD_SHEET,
                            int(selected["_row"]),
                            {
                                "unvan_ad_soyad": new_name,
                                "kurum": new_kurum,
                                "Gorev": new_gorev,
                                "Gun_ve_Saat": new_time,
                                "Salon": new_salon,
                                "Oturum_ID": new_oturum,
                                "Online": new_online,
                            },
                        )
                        st.cache_data.clear()
                        st.success("Moderatör kaydı güncellendi.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Güncelleme hatası: {e}")

        with mod_tab2:
            st.caption("İki satır arasında ya kişileri ya da oturum atamalarını değiştirebilirsin.")
            options = {row_label(r, ["unvan_ad_soyad", "Gorev", "Gun_ve_Saat", "Salon", "Oturum_ID"]): r for _, r in df_mod.iterrows()}
            c1, c2 = st.columns(2)
            left_label = c1.selectbox("1. kişi", list(options.keys()), key="swap_left")
            right_label = c2.selectbox("2. kişi", list(options.keys()), key="swap_right")
            swap_mode = st.radio("Değişim tipi", ["Oturumlarını değiştir", "Kişi isimlerini değiştir"], horizontal=True)
            if st.button("Yer Değiştir"):
                if not writable:
                    st.error("Yazma modu aktif değil.")
                elif left_label == right_label:
                    st.warning("İki farklı kişi seçmelisin.")
                else:
                    left = options[left_label]
                    right = options[right_label]
                    try:
                        if swap_mode == "Oturumlarını değiştir":
                            fields = ["Gun_ve_Saat", "Salon", "Oturum_ID", "Online"]
                        else:
                            fields = ["unvan_ad_soyad", "kurum"]
                        left_updates = {f: safe_str(right.get(f, "")) for f in fields if f in df_mod.columns}
                        right_updates = {f: safe_str(left.get(f, "")) for f in fields if f in df_mod.columns}
                        update_sheet_row(program_sheet_id, PROGRAM_MOD_SHEET, int(left["_row"]), left_updates)
                        update_sheet_row(program_sheet_id, PROGRAM_MOD_SHEET, int(right["_row"]), right_updates)
                        st.cache_data.clear()
                        st.success("Değişim yapıldı.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Değişim hatası: {e}")


with tab_matbaa:
    st.subheader("Matbaa Çıktıları")
    if df_master.empty:
        st.error("Matbaa için program ana verisi hazırlanamadı.")
    else:
        st.info("Bu bölüm önceki matbaa mantığını korur; güncel Sheet verisinden Excel ve Word dosyaları üretir.")
        c1, c2, c3 = st.columns(3)
        c1.metric("Toplam Program Kaydı", len(df_master))
        c2.metric("Yüz yüze", (df_master["Sunum_Tipi"] == "Yuzyuze").sum())
        c3.metric("Online", (df_master["Sunum_Tipi"] == "Online").sum())

        if st.button("Matbaayı Çalıştır ve Dosyaları Hazırla"):
            with st.spinner("Excel ve Word çıktıları hazırlanıyor..."):
                df_yz = df_master[df_master["Sunum_Tipi"] == "Yuzyuze"]
                df_on = df_master[df_master["Sunum_Tipi"] == "Online"]
                st.session_state["matbaa_outputs"] = {
                    "excel_yz": excel_bas(df_yz, raw_data.get("ozel", pd.DataFrame()), raw_data.get("mod", pd.DataFrame()), "YUZYUZE").getvalue(),
                    "word_yz": word_bas(df_yz, raw_data.get("ozel", pd.DataFrame()), raw_data.get("mod", pd.DataFrame()), "YUZYUZE").getvalue(),
                    "excel_on": excel_bas(df_on, raw_data.get("ozel", pd.DataFrame()), raw_data.get("mod", pd.DataFrame()), "ONLINE").getvalue(),
                    "word_on": word_bas(df_on, raw_data.get("ozel", pd.DataFrame()), raw_data.get("mod", pd.DataFrame()), "ONLINE").getvalue(),
                    "ozet": oturum_ozeti_olustur(df_master).getvalue(),
                }
                st.success("Matbaa çıktıları hazır.")

        outputs = st.session_state.get("matbaa_outputs")
        if outputs:
            c1, c2 = st.columns(2)
            with c1:
                st.write("#### Yüz Yüze Program")
                st.download_button("Excel İndir", outputs["excel_yz"], file_name="Nihai_Program_Yuzyuze.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                st.download_button("Word İndir", outputs["word_yz"], file_name="Nihai_Program_Yuzyuze.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            with c2:
                st.write("#### Online Program")
                st.download_button("Excel İndir", outputs["excel_on"], file_name="Nihai_Program_Online.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                st.download_button("Word İndir", outputs["word_on"], file_name="Nihai_Program_Online.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            st.download_button("Oturum Röntgen Özeti İndir", outputs["ozet"], file_name="Oturum_Rontgen_Ozeti.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


with tab_reports:
    st.subheader("Raporlar ve Excel İndirme Merkezi")
    c1, c2 = st.columns(2)
    c1.download_button("Genel Katılımcı Analizi", data=make_genel_report(katilimcilar), file_name="IHMC_Genel_Rapor.xlsx", use_container_width=True)
    c1.download_button("Sadece Online Katılımcılar", data=make_katilim_turu_report(katilimcilar, "ONLINE"), file_name="Online_Katilimcilar.xlsx", use_container_width=True)
    c2.download_button("Ödenmemiş Bildiriler", data=make_unpaid_report(bildiriler, katilimcilar), file_name="Odenmemis_Bildiriler.xlsx", use_container_width=True)
    c2.download_button("Sadece Fiziki Katılımcılar", data=make_katilim_turu_report(katilimcilar, "FİZİKİ"), file_name="Fiziki_Katilimcilar.xlsx", use_container_width=True)

    st.write("#### Yemek ve Etkinlik Listeleri")
    m1, m2, m3 = st.columns(3)
    m1.download_button("07 Mayıs Öğle", data=make_meal_report(katilimcilar, ["7", "OGLE"]), file_name="07_Mayis_Ogle.xlsx", use_container_width=True)
    m2.download_button("07 Mayıs Gala", data=make_meal_report(katilimcilar, ["GALA"]), file_name="07_Mayis_Gala.xlsx", use_container_width=True)
    m3.download_button("08 Mayıs Öğle", data=make_meal_report(katilimcilar, ["8", "OGLE"]), file_name="08_Mayis_Ogle.xlsx", use_container_width=True)

    st.write("#### Zenginleştirilmiş Program")
    df_p = raw_data.get("program_plan", pd.DataFrame()).copy()
    if df_p.empty:
        st.warning("Program tablosu boş.")
    else:
        enriched_rows = []
        for _, row in df_p.iterrows():
            row_dict = row.to_dict()
            row_str = super_temiz(" ".join([str(x) for x in row.values if pd.notna(x)]))
            b_adi = ""
            yazarlar = ""
            sunucu = ""
            for _, b_v in bildiriler.items():
                if super_temiz(b_v["Orijinal İsim"]) in row_str:
                    b_adi = turkce_buyuk(b_v["Orijinal İsim"])
                    yazarlar = ", ".join(
                        [
                            f"{turkce_buyuk(katilimcilar.get(y, {}).get('Unvan', ''))} {turkce_buyuk(katilimcilar.get(y, {}).get('Orijinal İsim', y))}".strip()
                            for y in b_v["Yazarlar"]
                        ]
                    )
                    sun = b_v["Sunucu"]
                    sunucu = f"{turkce_buyuk(katilimcilar.get(sun, {}).get('Unvan', ''))} {turkce_buyuk(katilimcilar.get(sun, {}).get('Orijinal İsim', sun))}".strip()
                    break
            row_dict["Eşleşen Bildiri Adı (Sistem)"] = b_adi
            row_dict["Yazarlar"] = yazarlar
            row_dict["Sunucu"] = sunucu
            enriched_rows.append(row_dict)

        df_enriched = pd.DataFrame(enriched_rows)
        filter_cols = st.columns(3)
        secili_gun = filter_cols[0].multiselect("Gün", sorted({gun_key(x) for x in df_p["Gun_ve_Saat"].dropna().astype(str)}) if "Gun_ve_Saat" in df_p.columns else [])
        secili_salon = filter_cols[1].multiselect("Salon", sorted(df_p["Salon"].dropna().unique()) if "Salon" in df_p.columns else [])
        secili_tip = filter_cols[2].multiselect("Sunum tipi", sorted(df_p["Sunum_Tipi"].dropna().unique()) if "Sunum_Tipi" in df_p.columns else [])

        df_filtered = df_enriched.copy()
        if "Gun_ve_Saat" in df_filtered.columns and secili_gun:
            df_filtered = df_filtered[df_filtered["Gun_ve_Saat"].apply(gun_key).isin(secili_gun)]
        if "Salon" in df_filtered.columns and secili_salon:
            df_filtered = df_filtered[df_filtered["Salon"].isin(secili_salon)]
        if "Sunum_Tipi" in df_filtered.columns and secili_tip:
            df_filtered = df_filtered[df_filtered["Sunum_Tipi"].isin(secili_tip)]

        out_prog = io.BytesIO()
        df_filtered.to_excel(out_prog, index=False)
        st.download_button(f"Programı İndir ({len(df_filtered)} kayıt)", data=out_prog.getvalue(), file_name="Bildiri_Sunum_Programi_Zengin.xlsx", use_container_width=True)


with tab_telegram:
    st.subheader("Telegram Bot Notu")
    st.info(
        "Telegram botunun uzun süreli polling ile Streamlit içinde çalışması sağlıklı değildir. "
        "Streamlit her etkileşimde script'i yeniden çalıştırdığı için bot döngüsü kilitlenebilir veya birden fazla bot kopyası açılabilir."
    )
    st.write(
        "Bu panel Telegram botunun kullandığı veri mantığını zaten içeriyor. Botu canlı tutmak istersen en temiz yapı: "
        "bu dosyadaki veri okuma/temizleme mantığını ayrı bir `core.py` dosyasına almak ve Telegram botunu ayrı bir worker olarak çalıştırmak."
    )
    st.warning("Bot token'ı koda gömülmemeli. BotFather'dan token'ı yenileyip ortam değişkeni veya secrets üzerinden kullanmalısın.")
